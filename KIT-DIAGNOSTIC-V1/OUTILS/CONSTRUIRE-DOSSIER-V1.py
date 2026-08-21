from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FORET = "183128"
ENCRE = "17221D"
PAPIER = "F5F1E8"
IVOIRE = "FBF9F4"
PIERRE = "D8D3C8"
SAUGE = "85938A"
CUIVRE = "A36F43"
BLANC = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=110, start=150, bottom=110, end=150):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=PIERRE, size=6):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        node = borders.find(qn(tag))
        if node is None:
            node = OxmlElement(tag)
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def set_font(run, name="Arial", size=10.5, color=ENCRE, bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.32)
    section.footer_distance = Inches(0.34)

    background = OxmlElement("w:background")
    background.set(qn("w:color"), PAPIER)
    doc._element.insert(0, background)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(ENCRE)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.16

    for style_name, size, before, after in (
        ("Heading 1", 22, 12, 8),
        ("Heading 2", 15, 10, 5),
        ("Heading 3", 11, 8, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Georgia"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Georgia")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Georgia")
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = rgb(FORET)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(10.5)
        style.font.color.rgb = rgb(ENCRE)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.16

    for current in doc.sections:
        set_running_header_footer(current)


def set_running_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("R S P I R   |   DOSSIER DE DÉCISION")
    set_font(run, size=7.5, color=SAUGE, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("CONFIDENTIEL   •   MODÈLE V1")
    set_font(run, size=7.5, color=SAUGE)


def add_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(text.upper())
    set_font(run, size=8, color=CUIVRE, bold=True)
    run.font.all_caps = True
    return p


def add_title(doc, text, size=30, color=FORET, after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    set_font(run, name="Georgia", size=size, color=color)
    return p


def add_body(doc, text, bold_lead=None, color=ENCRE, italic=False, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_font(first, size=10.5, color=color, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_font(rest, size=10.5, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_font(run, size=10.5, color=color, italic=italic)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    set_font(run, size=10.5, color=ENCRE)
    return p


def add_callout(doc, label, text, fill=IVOIRE, accent=CUIVRE):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=PIERRE, size=5)
    set_cell_margins(cell, top=180, start=210, bottom=180, end=210)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label.upper())
    set_font(r, size=8, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_font(r2, name="Georgia", size=13, color=FORET)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_section_page(doc, number, title, anchor):
    doc.add_page_break()
    add_label(doc, f"{number:02d}  •  Dossier de décision")
    add_title(doc, title, size=23)
    add_body(doc, anchor, color=SAUGE, italic=True, after=14)


def add_placeholder(doc, label, instruction, size="court"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(label.upper())
    set_font(r, size=8, color=CUIVRE, bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(8)
    r2 = p2.add_run(f"[À compléter — {instruction}. Format : {size}.]")
    set_font(r2, size=10.5, color=SAUGE, italic=True)


def add_two_column_cards(doc, cards):
    rows = (len(cards) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    set_table_widths(table, [4680, 4680])
    for index, (label, text) in enumerate(cards):
        cell = table.rows[index // 2].cells[index % 2]
        set_cell_shading(cell, IVOIRE)
        set_cell_border(cell, PIERRE, 5)
        set_cell_margins(cell, 150, 170, 150, 170)
        p = cell.paragraphs[0]
        r = p.add_run(label.upper())
        set_font(r, size=7.5, color=CUIVRE, bold=True)
        p2 = cell.add_paragraph()
        r2 = p2.add_run(text)
        set_font(r2, size=9.5, color=SAUGE, italic=True)
    if len(cards) % 2:
        cell = table.rows[-1].cells[-1]
        set_cell_shading(cell, PAPIER)
        set_cell_border(cell, PAPIER, 1)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, value in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, FORET)
        set_cell_border(cell, FORET, 5)
        set_cell_margins(cell, 120, 140, 120, 140)
        p = cell.paragraphs[0]
        r = p.add_run(value.upper())
        set_font(r, size=7.5, color=PAPIER, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            cell = cells[idx]
            set_cell_shading(cell, IVOIRE)
            set_cell_border(cell, PIERRE, 5)
            set_cell_margins(cell, 120, 140, 120, 140)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            set_font(r, size=8.8, color=SAUGE if value.startswith("[") else ENCRE, italic=value.startswith("["))
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def build(output_path):
    doc = Document()
    configure_document(doc)

    # Couverture
    add_label(doc, "Diagnostic RSPIR")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(84)
    add_title(doc, "Dossier de décision", size=35, after=12)
    add_body(doc, "Une meilleure prochaine décision, suffisamment étayée pour savoir pourquoi elle passe en premier, ce qui doit être prêt et dans quel ordre avancer.", color=SAUGE, italic=True, after=28)
    add_callout(doc, "Client", "[Nom de l’organisation]", fill=IVOIRE)
    add_two_column_cards(doc, [
        ("Question de départ", "[Question ayant déclenché le Diagnostic]"),
        ("Date", "[AAAA-MM-JJ]"),
        ("Préparé par", "RSPIR"),
        ("Statut", "[Version validée / À valider]"),
    ])
    add_body(doc, "Document confidentiel. Le client demeure libre d’exécuter la décision avec son équipe, RSPIR ou un autre partenaire.", color=SAUGE, italic=True, after=0)

    # Lecture rapide
    doc.add_page_break()
    add_label(doc, "Lecture rapide")
    add_title(doc, "La décision en bref", size=27)
    add_callout(doc, "Prochaine décision", "[Décision ou action précise à prendre]", fill=IVOIRE)
    add_two_column_cards(doc, [
        ("Pourquoi elle passe en premier", "[Une phrase fondée sur la comparaison]"),
        ("Niveau de confiance", "[Élevé / moyen / faible + raison]"),
        ("Condition critique", "[Ce qui doit être vrai avant d’avancer]"),
        ("À ne pas faire maintenant", "[Élément principal de la Stop List]"),
    ])
    add_placeholder(doc, "Premier mouvement", "propriétaire, action, preuve attendue et moment de revue", "quatre lignes maximum")
    add_body(doc, "Cette page résume le Dossier. Elle ne remplace ni les preuves ni les incertitudes présentées plus loin.", color=SAUGE, italic=True)

    add_section_page(doc, 1, "Contexte et question de départ", "Pourquoi le Diagnostic a été lancé.")
    add_placeholder(doc, "Déclencheur", "événement, friction ou décision imminente", "trois phrases maximum")
    add_placeholder(doc, "Résultat recherché", "ce que le dirigeant veut améliorer, sans confondre destination économique et promesse du Diagnostic", "deux phrases")
    add_placeholder(doc, "Question de départ", "formulation confirmée qui guidera le Diagnostic", "une phrase")
    add_two_column_cards(doc, [
        ("Décisions bloquées", "[Décisions que le client n’arrive pas à prendre]"),
        ("Contraintes connues", "[Budget, délai, capacité, risque ou contexte]"),
    ])

    add_section_page(doc, 2, "Lecture organisationnelle", "Voir suffisamment large pour éviter d’approfondir le mauvais problème.")
    add_placeholder(doc, "Vue synthétique", "les trois à cinq signaux qui peuvent réellement déplacer le choix du périmètre", "cinq points maximum")
    add_matrix(doc, ["Signal", "Preuve", "Effet possible"], [
        ("[Signal important]", "[Fait observé + source/date]", "[Comment il influence la priorité]"),
        ("[Signal important]", "[Fait observé + source/date]", "[Comment il influence la priorité]"),
        ("[Signal important]", "[Fait observé + source/date]", "[Comment il influence la priorité]"),
    ], [2500, 3600, 3260])
    add_body(doc, "Les 31 dimensions ont servi de radar. Elles ne constituent pas 31 audits.", color=SAUGE, italic=True)

    add_section_page(doc, 3, "Priorités candidates", "Comparer les zones significatives avant de choisir.")
    add_matrix(doc, ["Candidate", "Valeur", "Faisabilité", "Préparation", "Risque / délai", "Confiance"], [
        ("[Zone A]", "[court]", "[court]", "[court]", "[court]", "[niveau]"),
        ("[Zone B]", "[court]", "[court]", "[court]", "[court]", "[niveau]"),
        ("[Zone C]", "[court]", "[court]", "[court]", "[court]", "[niveau]"),
    ], [1800, 1500, 1500, 1500, 1560, 1500])
    add_placeholder(doc, "Jugement", "ce que le score éclaire et ce qu’il ne décide pas", "un court paragraphe")
    add_callout(doc, "Sortie valide", "[Priorité claire / choix conditionnel / incertitude à réduire]", fill=IVOIRE)

    add_section_page(doc, 4, "Priorité recommandée", "Ce qui passe en premier — et pourquoi les alternatives passent après.")
    add_callout(doc, "Priorité", "[Périmètre vertical recommandé]", fill=IVOIRE)
    add_placeholder(doc, "Raisonnement", "lien entre la preuve, la valeur, la faisabilité, la préparation, le risque et le délai", "cinq à sept lignes")
    add_matrix(doc, ["Alternative", "Pourquoi elle ne passe pas en premier", "Quand la revoir"], [
        ("[Alternative 1]", "[raison concise]", "[déclencheur]"),
        ("[Alternative 2]", "[raison concise]", "[déclencheur]"),
    ], [2600, 4300, 2460])
    add_placeholder(doc, "Ce qui pourrait invalider le choix", "information déterminante encore fragile ou inconnue", "trois points maximum")

    add_section_page(doc, 5, "Analyse du périmètre", "Descendre jusqu’au niveau nécessaire pour décider, pas pour construire.")
    add_two_column_cards(doc, [
        ("Humains", "[propriété, exécution, validation, dépendances]"),
        ("Processus", "[flux réel, exceptions, reprises, goulot]"),
        ("Logiciels", "[systèmes pertinents, ruptures, capacités existantes]"),
        ("Données et gouvernance", "[existence, accès, cohérence, risque proportionnel]"),
    ])
    add_placeholder(doc, "IA, si pertinente", "capacité ajoutée, autonomie, validation humaine, erreur acceptable et permissions", "un bloc court — sinon « non pertinente maintenant »")
    add_placeholder(doc, "Friction dominante", "où le résultat se brise réellement", "une phrase")

    add_section_page(doc, 6, "Cause et mécanisme", "Expliquer ce qui semble produire le problème, sans transformer une hypothèse en fait.")
    add_callout(doc, "Mécanisme principal", "[Cause plausible → mécanisme → effet observé]", fill=IVOIRE)
    add_matrix(doc, ["Élément", "Statut", "Fondement"], [
        ("[Élément 1]", "Fait", "[observation + source/date]"),
        ("[Élément 2]", "Hypothèse", "[indice + niveau de confiance]"),
        ("[Élément 3]", "Inconnue", "[preuve manquante]"),
    ], [2600, 1700, 5060])
    add_placeholder(doc, "Interventions raisonnables", "humain, processus, logiciel, intégration, automatisation, IA ou combinaison", "trois options maximum")

    add_section_page(doc, 7, "Cas économique", "Quantifier ce qui se défend. Qualifier le reste.")
    add_callout(doc, "Niveau de preuve économique", "[A — quantifiable / B — estimable / C — stratégique ou qualitatif]", fill=IVOIRE)
    add_matrix(doc, ["Facteur", "Situation actuelle", "Effet plausible", "Source / confiance"], [
        ("[Temps, coût, volume ou risque]", "[valeur ou inconnue]", "[fourchette ou effet]", "[source/date/niveau]"),
        ("[Capacité, qualité ou dépendance]", "[constat]", "[effet]", "[source/date/niveau]"),
    ], [2400, 2200, 2500, 2260])
    add_placeholder(doc, "Coût de l’inaction", "ce qui se poursuit ou s’aggrave si rien ne change", "un court paragraphe")
    add_body(doc, "Aucun rendement précis n’est garanti. Les hypothèses demeurent visibles.", color=SAUGE, italic=True)

    add_section_page(doc, 8, "Conditions de préparation", "Ce qui doit être vrai dans le périmètre avant ou pendant l’intervention.")
    add_matrix(doc, ["Condition", "État", "Écart", "Responsable"], [
        ("[Condition critique]", "[présente / partielle / absente]", "[ce qui manque]", "[rôle]"),
        ("[Condition importante]", "[état]", "[écart]", "[rôle]"),
        ("[Condition utile]", "[état]", "[écart]", "[rôle]"),
    ], [2900, 1900, 2860, 1700])
    add_placeholder(doc, "Seuil avant d’avancer", "condition minimale permettant la prochaine décision", "une phrase")
    add_body(doc, "Ces conditions concernent la priorité recommandée; elles ne constituent pas un score global de maturité de l’entreprise.", color=SAUGE, italic=True)

    add_section_page(doc, 9, "Séquence", "Maintenant, ensuite, plus tard — ou pas maintenant.")
    add_matrix(doc, ["Moment", "Décision ou mouvement", "Preuve de passage"], [
        ("Maintenant", "[première décision]", "[preuve observable]"),
        ("Ensuite", "[mouvement suivant]", "[condition d’entrée]"),
        ("Plus tard", "[élément différé]", "[déclencheur de réévaluation]"),
    ], [1800, 4300, 3260])
    add_placeholder(doc, "Première valeur", "premier effet utile attendu, sans promesse artificielle de délai", "deux phrases")
    add_body(doc, "La conception détaillée et l’implantation appartiennent à des phases distinctes.", color=SAUGE, italic=True)

    add_section_page(doc, 10, "Stop List", "Éviter une mauvaise dépense est une forme réelle de rendement.")
    add_matrix(doc, ["À éviter ou reporter", "Pourquoi", "Condition de réouverture"], [
        ("[Initiative 1]", "[risque, prématurité ou faible valeur]", "[preuve ou changement requis]"),
        ("[Initiative 2]", "[raison]", "[déclencheur]"),
        ("[Initiative 3]", "[raison]", "[déclencheur]"),
    ], [3000, 3500, 2860])
    add_callout(doc, "Règle", "Ne pas automatiser, acheter ou construire simplement parce que l’option existe.", fill=IVOIRE)

    add_section_page(doc, 11, "Incertitudes", "Rendre visible ce qui n’est pas suffisamment connu et ce que cela change.")
    add_matrix(doc, ["Incertitude", "Effet sur la décision", "Comment la réduire", "Responsable"], [
        ("[Inconnue déterminante]", "[ce qu’elle pourrait changer]", "[mesure, observation ou preuve]", "[rôle]"),
        ("[Hypothèse fragile]", "[effet]", "[test proportionné]", "[rôle]"),
    ], [2500, 2500, 2760, 1600])
    add_body(doc, "« Nous ne savons pas encore » est une conclusion valide lorsqu’elle mène à une mesure précise avant d’investir.", color=SAUGE, italic=True)
    add_placeholder(doc, "Limite de confiance", "ce que le Dossier ne permet pas encore d’affirmer", "deux phrases")

    add_section_page(doc, 12, "Prochaine décision", "Une action assez précise pour que le dirigeant sache quoi décider, avec qui et sur quelle preuve.")
    add_callout(doc, "Décision", "[Décision ou action exacte]", fill=IVOIRE)
    add_two_column_cards(doc, [
        ("Propriétaire", "[Personne responsable de décider ou d’agir]"),
        ("Moment", "[Date ou déclencheur externe, sinon à convenir]"),
        ("Entrée nécessaire", "[Preuve ou condition minimale]"),
        ("Sortie attendue", "[Décision, validation ou résultat observable]"),
    ])
    add_placeholder(doc, "Mandat suivant possible", "décrire seulement la nature de la phase 2, sans architecture ni sélection détaillée", "trois lignes maximum")
    add_body(doc, "Le client peut exécuter cette décision avec son équipe, RSPIR ou un autre partenaire.", color=SAUGE, italic=True)

    # Dernière page de contrôle client
    doc.add_page_break()
    add_label(doc, "Appropriation")
    add_title(doc, "La décision tient-elle?", size=27)
    add_body(doc, "Le Diagnostic est terminé lorsque la décision est suffisamment étayée — pas simplement parce que le rapport existe.", color=SAUGE, italic=True, after=14)
    add_matrix(doc, ["Critère", "Question de contrôle"], [
        ("Direction", "La priorité ou les scénarios sont-ils clairs?"),
        ("Justification", "Pourquoi passent-ils devant les alternatives?"),
        ("Preuve", "Les observations et hypothèses importantes sont-elles identifiables?"),
        ("Conditions", "Les prérequis et obstacles sont-ils connus?"),
        ("Économie", "La valeur est-elle établie au niveau de précision raisonnable?"),
        ("Séquence", "Savons-nous ce qui vient maintenant, ensuite et plus tard?"),
        ("Action", "La prochaine décision est-elle suffisamment précise?"),
    ], [2200, 7160])
    add_callout(doc, "Transfert de compréhension", "Le dirigeant doit pouvoir expliquer la logique à son équipe sans dépendre de RSPIR pour traduire son propre dossier.", fill=IVOIRE)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 2:
        raise SystemExit("Usage: CONSTRUIRE-DOSSIER-V1.py <sortie.docx>")
    build(Path(sys.argv[1]).resolve())
